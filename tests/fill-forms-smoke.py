#!/usr/bin/env python3
"""Local dependency-backed smoke test for the packaged fill-forms skill."""

from __future__ import annotations

import hashlib
import json
import stat
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "skills" / "fill-forms"
SCRIPTS = SKILL / "scripts"


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def run(script: str, *args: object, expected: int = 0) -> subprocess.CompletedProcess:
    result = subprocess.run(
        [sys.executable, str(SCRIPTS / script), *(str(value) for value in args)],
        text=True,
        capture_output=True,
        check=False,
    )
    if result.returncode != expected:
        raise AssertionError(
            f"{script} returned {result.returncode}, expected {expected}\n"
            f"stdout:\n{result.stdout}\nstderr:\n{result.stderr}"
        )
    return result


def pdf_smoke(folder: Path) -> None:
    source = folder / "blank.pdf"
    completed = folder / "completed.pdf"
    config = folder / "form.json"

    document = fitz.open()
    page = document.new_page()
    page.insert_text((50, 80), "Name:")
    page.insert_text((50, 130), "Consent:")
    widget = fitz.Widget()
    widget.field_name = "legal_name"
    widget.field_label = "Legal name"
    widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
    widget.rect = fitz.Rect(110, 55, 300, 85)
    page.add_widget(widget)
    checkbox = fitz.Widget()
    checkbox.field_name = "accepted"
    checkbox.field_label = "Accepted"
    checkbox.field_type = fitz.PDF_WIDGET_TYPE_CHECKBOX
    checkbox.rect = fitz.Rect(110, 105, 125, 120)
    page.add_widget(checkbox)
    document.save(source)
    document.close()
    source_hash = digest(source)

    config.write_text(
        json.dumps(
            {
                "src": source.name,
                "out": completed.name,
                "mode": "both",
                "acroform": {
                    "legal_name": "Example Person",
                    "accepted": True,
                },
                "overlay": [
                    {"page": 0, "x": 110, "y": 130, "check": True},
                ],
            }
        ),
        encoding="utf-8",
    )

    run("inspect_pdf.py", source)
    grid_refused = run("inspect_pdf.py", source, expected=1)
    assert "Grid output already exists" in grid_refused.stderr
    run("fill_pdf.py", config)
    assert source_hash == digest(source)
    assert completed.is_file()
    grid_dir = folder / "blank_grid"
    preview_dir = folder / "completed_preview"
    grid_image = grid_dir / "page-1.png"
    preview_image = preview_dir / "page-1.png"
    assert grid_image.is_file()
    assert preview_image.is_file()
    assert stat.S_IMODE(completed.stat().st_mode) == 0o600
    assert stat.S_IMODE(grid_dir.stat().st_mode) == 0o700
    assert stat.S_IMODE(preview_dir.stat().st_mode) == 0o700
    assert stat.S_IMODE(grid_image.stat().st_mode) == 0o600
    assert stat.S_IMODE(preview_image.stat().st_mode) == 0o600

    output = fitz.open(completed)
    try:
        text = output[0].get_text()
        assert "Example Person" in text
        assert "X" in text
        widgets = list(output[0].widgets() or [])
        values = {widget.field_name: widget.field_value for widget in widgets}
        assert values["legal_name"] == "Example Person"
        assert values["accepted"] not in (False, None, "Off", "No")
    finally:
        output.close()

    refused = run("fill_pdf.py", config, expected=1)
    assert "Output already exists" in refused.stderr

    unknown = folder / "unknown-field.json"
    unknown.write_text(
        json.dumps(
            {
                "src": source.name,
                "out": "unknown-field.pdf",
                "mode": "acroform",
                "acroform": {"missing-field": "value"},
            }
        ),
        encoding="utf-8",
    )
    rejected = run("fill_pdf.py", unknown, expected=1)
    assert "Unknown AcroForm fields" in rejected.stderr
    assert not (folder / "unknown-field.pdf").exists()


def docx_smoke(folder: Path) -> None:
    source = folder / "template.docx"
    completed = folder / "completed.docx"
    data = folder / "docx-data.json"

    document = Document()
    document.add_paragraph("Applicant: {{name}}")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Reference: {{reference}}"
    document.save(source)
    source_hash = digest(source)
    data.write_text(
        json.dumps({"name": "Example Person", "reference": "EXAMPLE-001"}),
        encoding="utf-8",
    )

    run("dump_docx.py", source)
    run("fill_docx.py", source, completed, "--data", data)
    assert source_hash == digest(source)
    assert stat.S_IMODE(completed.stat().st_mode) == 0o600

    output = Document(completed)
    assert "Example Person" in "\n".join(p.text for p in output.paragraphs)
    assert output.tables[0].cell(0, 0).text == "Reference: EXAMPLE-001"

    unused_data = folder / "unused-data.json"
    unused_output = folder / "unused.docx"
    unused_data.write_text(
        json.dumps({"not_in_template": "value"}),
        encoding="utf-8",
    )
    rejected = run(
        "fill_docx.py",
        source,
        unused_output,
        "--data",
        unused_data,
        expected=1,
    )
    assert "Unused placeholders" in rejected.stderr
    assert not unused_output.exists()


def signature_smoke(folder: Path) -> None:
    source = folder / "signature-source.png"
    completed = folder / "signature-prepared.png"
    image = Image.new("RGB", (400, 120), "white")
    draw = ImageDraw.Draw(image)
    draw.line([(45, 75), (130, 30), (220, 82), (350, 45)], fill="black", width=7)
    image.save(source)
    source_hash = digest(source)

    run("prepare_signature.py", source, completed)
    assert source_hash == digest(source)
    assert stat.S_IMODE(completed.stat().st_mode) == 0o600
    prepared = Image.open(completed)
    try:
        assert prepared.mode == "RGBA"
        assert prepared.getchannel("A").getbbox() is not None
    finally:
        prepared.close()


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="nc-fill-forms-smoke-") as temporary:
        folder = Path(temporary)
        pdf_smoke(folder)
        docx_smoke(folder)
        signature_smoke(folder)
    print("fill-forms smoke: PASS")


if __name__ == "__main__":
    main()
