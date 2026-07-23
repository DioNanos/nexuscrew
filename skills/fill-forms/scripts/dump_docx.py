#!/usr/bin/env python3
"""Print paragraphs, tables, headers and footers from a DOCX template."""

from __future__ import annotations

import argparse
from pathlib import Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Inspect text and table structure in a local DOCX form."
    )
    parser.add_argument("docx", type=Path)
    return parser.parse_args()


def load_docx():
    try:
        import docx
    except ImportError as exc:
        raise SystemExit(
            "Missing dependency: python-docx. With user consent, install the "
            "packages listed in the skill's requirements.txt."
        ) from exc
    return docx


def print_paragraphs(paragraphs, heading: str) -> None:
    print(f"== {heading} ==")
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"[{index}] {text}")


def print_tables(tables, heading: str) -> None:
    for table_index, table in enumerate(tables):
        print(
            f"\n== {heading} TABLE {table_index} "
            f"({len(table.rows)}x{len(table.columns)}) =="
        )
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  row {row_index}: {cells}")


def main() -> None:
    args = parse_args()
    if not args.docx.is_file():
        raise SystemExit(f"DOCX not found or not a regular file: {args.docx}")

    docx = load_docx()
    try:
        document = docx.Document(args.docx)
    except Exception as exc:
        raise SystemExit(f"Could not open DOCX: {exc}") from exc

    print_paragraphs(document.paragraphs, "BODY PARAGRAPHS")
    print_tables(document.tables, "BODY")

    for section_index, section in enumerate(document.sections):
        print_paragraphs(section.header.paragraphs, f"HEADER {section_index}")
        print_tables(section.header.tables, f"HEADER {section_index}")
        print_paragraphs(section.footer.paragraphs, f"FOOTER {section_index}")
        print_tables(section.footer.tables, f"FOOTER {section_index}")


if __name__ == "__main__":
    main()
