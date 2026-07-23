#!/usr/bin/env python3
"""Fill explicit placeholders in a local DOCX template."""

from __future__ import annotations

import argparse
import json
import os
import re
import tempfile
from collections import Counter
from pathlib import Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Replace {{placeholders}} in a DOCX form using a JSON mapping."
    )
    parser.add_argument("src", type=Path)
    parser.add_argument("out", type=Path)
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument(
        "--literal-keys",
        action="store_true",
        help="also replace unbraced mapping keys; use only for an inspected template",
    )
    parser.add_argument(
        "--allow-unused",
        action="store_true",
        help="allow mapping keys that were not found in the document",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="replace an existing output file, never the source",
    )
    return parser.parse_args()


def load_docx():
    try:
        import docx
    except ImportError as exc:
        raise SystemExit(
            "Missing dependency: python-docx. With user consent, install the "
            "packages listed in the skill's requirements.txt."
        ) from exc
    return docx


def load_mapping(path: Path) -> dict[str, str]:
    if not path.is_file():
        raise SystemExit(f"Data file not found: {path}")
    if path.stat().st_size > 2 * 1024 * 1024:
        raise SystemExit("Data JSON exceeds the 2 MiB safety limit.")

    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise SystemExit(f"Could not read data JSON: {exc}") from exc

    if not isinstance(raw, dict) or not raw:
        raise SystemExit("Data JSON must be a non-empty object.")

    mapping: dict[str, str] = {}
    for key, value in raw.items():
        if (
            not isinstance(key, str)
            or not key.strip()
            or len(key) > 128
            or "{{" in key
            or "}}" in key
        ):
            raise SystemExit("Every placeholder key must be a plain non-empty string.")
        if value is None:
            mapping[key] = ""
        elif isinstance(value, (str, int, float, bool)):
            mapping[key] = str(value)
        else:
            raise SystemExit(f"Placeholder {key!r} has a non-scalar value.")
    return mapping


def token_pattern(mapping: dict[str, str], literal_keys: bool):
    token_to_key = {}
    for key in mapping:
        token_to_key[f"{{{{{key}}}}}"] = key
        if literal_keys:
            token_to_key[key] = key
    ordered = sorted(token_to_key, key=len, reverse=True)
    return re.compile("|".join(re.escape(token) for token in ordered)), token_to_key


def replace_in_paragraph(paragraph, pattern, token_to_key, mapping, counts) -> None:
    original = "".join(run.text for run in paragraph.runs)
    if not original:
        return

    def replacement(match):
        key = token_to_key[match.group(0)]
        counts[key] += 1
        return mapping[key]

    updated = pattern.sub(replacement, original)
    if updated == original:
        return
    if not paragraph.runs:
        paragraph.add_run(updated)
        return
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def process_tables(tables, process_paragraph) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                process_tables(cell.tables, process_paragraph)


def save_atomic(document, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    handle = tempfile.NamedTemporaryFile(
        prefix=f".{output.name}.",
        suffix=".tmp",
        dir=output.parent,
        delete=False,
    )
    temporary = Path(handle.name)
    handle.close()
    try:
        document.save(temporary)
        os.replace(temporary, output)
        os.chmod(output, 0o600)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise


def main() -> None:
    args = parse_args()
    source = args.src.resolve()
    output_unresolved = args.out.expanduser()
    if output_unresolved.is_symlink():
        raise SystemExit("Output must not be a symbolic link.")
    output = output_unresolved.resolve()
    data_path = args.data.resolve()

    if not source.is_file():
        raise SystemExit(f"DOCX not found or not a regular file: {source}")
    if source == output:
        raise SystemExit("Source and output must be different files.")
    if output.exists() and not args.overwrite:
        raise SystemExit("Output already exists. Use --overwrite for an intentional revision.")

    mapping = load_mapping(data_path)
    pattern, token_to_key = token_pattern(mapping, args.literal_keys)
    counts = Counter()
    docx = load_docx()

    try:
        document = docx.Document(source)
    except Exception as exc:
        raise SystemExit(f"Could not open DOCX: {exc}") from exc

    def process(paragraph):
        replace_in_paragraph(
            paragraph,
            pattern,
            token_to_key,
            mapping,
            counts,
        )

    for paragraph in document.paragraphs:
        process(paragraph)
    process_tables(document.tables, process)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            process(paragraph)
        process_tables(section.header.tables, process)
        for paragraph in section.footer.paragraphs:
            process(paragraph)
        process_tables(section.footer.tables, process)

    unused = sorted(key for key in mapping if counts[key] == 0)
    if unused and not args.allow_unused:
        raise SystemExit(
            "Unused placeholders: "
            + ", ".join(repr(key) for key in unused)
            + ". Inspect the template or use --allow-unused deliberately."
        )

    save_atomic(document, output)
    print(f"Saved: {output}")
    print(f"Replacements: {sum(counts.values())}")
    if unused:
        print("Unused placeholders allowed: " + ", ".join(repr(key) for key in unused))
    print("Review the completed DOCX or export a local copy to PDF before handoff.")


if __name__ == "__main__":
    main()
